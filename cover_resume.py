from docx import Document
import comtypes.client
import os

def create_custom_cover_letter():
    letter = Document()

    # Giriş
    letter.add_paragraph("Dear Hiring Manager,\n")

    # 1. Paragraf – Tanıtım ve pozisyon ilgisi
    letter.add_paragraph(
        "I am writing to express my interest in the Frontend Developer position at your company. "
        "Based in Baku, Azerbaijan, I bring over 3 years of hands-on experience in building modern, scalable, "
        "and SEO-optimized web applications. My expertise lies primarily in Nuxt.js 3 and Vue.js 3, "
        "along with a strong understanding of SSR, CSR, and SSG architectures.\n"
    )

    # 2. Paragraf – Deneyim ve beceriler
    letter.add_paragraph(
        "Throughout my career, I have developed responsive and accessible web solutions, focusing on performance optimization "
        "and seamless user experiences. I have practical experience working with frontend tools such as Tailwind CSS, Pinia, and TypeScript, "
        "as well as backend technologies like Node.js, Express, and various databases, which helps me build well-rounded applications.\n"
    )

    # 3. Paragraf – Projeler ve motivasyon
    letter.add_paragraph(
        "Among my projects, I have created an SEO-optimized movie platform inspired by Netflix and a modern e-commerce SPA, "
        "both showcasing my skills in Vue.js, Nuxt.js, and CI/CD deployments. As a self-taught developer, I am passionate about continuous learning "
        "and staying up to date with the latest web technologies.\n"
    )

    # Kapanış ve teşekkür
    letter.add_paragraph(
        "I am excited about the opportunity to contribute to your team and bring my skills and enthusiasm to your projects. "
        "Thank you for considering my application.\n"
    )

    # İmza
    letter.add_paragraph("Best regards,\nFuad Aghayev\nBaku, Azerbaijan\n+994 055 700 14 02\nfuad0000010@gmail.com\nhttps://fuad-foliou.netlify.app/")

    # Dosya adı
    filename = "Fuad_Aghayev_CoverLetter.docx"
    letter.save(filename)
    print(f"Cover letter DOCX dosyası oluşturuldu: {filename}")
    return filename

def convert_docx_to_pdf(docx_path, pdf_path):
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(docx_path)
    doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
    doc.Close()
    word.Quit()
    print(f"PDF dosyası başarıyla oluşturuldu: {pdf_path}")

if __name__ == "__main__":
    docx_file = create_custom_cover_letter()
    pdf_file = docx_file.replace(".docx", ".pdf")

    convert_docx_to_pdf(os.path.abspath(docx_file), os.path.abspath(pdf_file))
