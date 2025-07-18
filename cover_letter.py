
from docx import Document
from pathlib import Path

#*  Başvuru formunda “upload additional file” gibi ikinci kutu varsa → Cover Letter PDF olarak yükle
#*“Comments” ya da “Tell us more” gibi boş kutular varsa → cover letter’ın özetini yapıştır
#*   APPLY dan sonra secenekler PDF YUKLE VE ISLEMLER GIR  ve  COMMENT VE TELL US MORE -YA dahil ede bilersin

def create_custom_cover_letter(company_name="Your Company", position_title="Remote Nuxt.js Developer"):
    letter = Document()

    # Giriş
    letter.add_paragraph("Dear Hiring Manager,\n")

    # 1. Paragraf – Kendini tanıt
    letter.add_paragraph(
        f"I am a Frontend Developer based in Baku, specializing in Nuxt.js 3, Vue.js, and modern JavaScript technologies. "
        f"I am writing to express my interest in the {position_title} position at {company_name}.\n"
    )

    # 2. Paragraf – Deneyim
    letter.add_paragraph(
        "I have hands-on experience developing responsive and dynamic web interfaces using Vue 3, Pinia, Tailwind CSS, "
        "and GitHub for version control. My recent internship allowed me to build real-world applications focused on "
        "performance, usability, and scalability.\n"
    )

    # 3. Paragraf – Kapanış
    letter.add_paragraph(
        "I would be excited to bring my enthusiasm and skills to your development team. "
        "Thank you for considering my application.\n"
    )

    # İmza
    letter.add_paragraph("Best regards,\nFuad Agayev")

    # Dosya adı
    filename = f"Fuad_Agayev_CoverLetter_{company_name.replace(' ', '_')}.docx"
    letter.save(filename)
    return filename

# Örnek: Vercel şirketine başvuru
create_custom_cover_letter(company_name="Vercel", position_title="Remote Frontend Engineer")
