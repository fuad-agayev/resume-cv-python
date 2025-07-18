from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# CV içeriğini yapılandır
cv = Document()

# Başlık - İsim
name = cv.add_heading("Fuad Agayev", 0)
name.alignment = WD_ALIGN_PARAGRAPH.CENTER

# İletişim bilgileri
contact = cv.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.add_run("Baku, Azerbaijan | +994 55 700 14 02 | fuad0000010@gmail.com | https://www.github.com/fuad-agayev\n").italic = True

# PROFİL (SUMMARY)
cv.add_heading("Summary", level=1)
cv.add_paragraph(
    "Frontend developer passionate about creating fast, user-friendly web interfaces. "
    "Skilled in Vue.js, Nuxt.js 3, and modern frontend technologies. Focused on remote opportunities "
    "as a Nuxt.js Developer. Eager to contribute clean, scalable, and maintainable code in collaborative teams."
)

# TEKNİK YETENEKLER (SKILLS)
cv.add_heading("Skills", level=1)
cv.add_paragraph(
    "- HTML5, CSS3, JavaScript (ES6+), TypeScript\n"
    "- Vue.js 3, Nuxt.js 3, Pinia, Composition API\n"
    "- Tailwind CSS, Sass, Responsive Design\n"
    "- Git, GitHub, REST APIs, SSR/CSR/SSG\n"
)

# DENEYİM (EXPERIENCE)
cv.add_heading("Experience", level=1)
cv.add_paragraph("NuxtJS Developer Intern – Baku, Azerbaijan (05/2022 – Present)\n"
    "- Developed responsive websites using HTML, CSS, and JavaScript\n"
    "- Built reusable Vue.js/Nuxt.js components with Composition API\n"
    "- Used GitHub for version control and collaborative development"
)

# KAYDET
cv.save("Fuad_Agayev_Resume_CV.docx")
print("CV başarıyla oluşturuldu: Fuad_Agayev_CV.docx")

#*  PDF -ye dönüştürme işlemi
import comtypes.client
import os

word = comtypes.client.CreateObject('Word.Application')

# Dosyanın tam yolunu al
doc_path = os.path.abspath("Fuad_Agayev_Resume_CV.docx")
pdf_path = os.path.abspath("Fuad_Agayev_Resume_CV.pdf")

# Word ile aç ve PDF'e kaydet
doc = word.Documents.Open(doc_path)
doc.SaveAs(pdf_path, FileFormat=17)
doc.Close()
word.Quit()

print("PDF başarıyla oluşturuldu:", pdf_path)