from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import comtypes.client
import os

def create_cv_docx(filename="Fuad_Aghayev_Resume_CV.docx"):
    cv = Document()

    # Font ayarı (Calibri 11 pt)
    style = cv.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # İsim - Başlık (level=0, ortalanmış)
    name = cv.add_heading('Fuad Aghayev', level=0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # İletişim bilgileri - Ortalanmış, italik, her bilgi ayrı satırda
    contact = cv.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("Baku, Azerbaijan | +994 055 700 14 02 | fuad0000010@gmail.com\n").italic = True
    contact.add_run("https://github.com/fuad-agayev\n").italic = True
    contact.add_run("https://www.linkedin.com/in/fuad-agayev-12b854328/\n").italic = True
    contact.add_run("https://fuad-foliou.netlify.app/").italic = True

    cv.add_paragraph()  # boşluk

    # Professional Summary
    cv.add_heading('Professional Summary', level=1)
    cv.add_paragraph(
        "Frontend Developer with 3 years of hands-on experience in building modern, scalable, and SEO-optimized web applications — "
        "primarily focused on Nuxt.js 3 development, along with Vue.js 3. Skilled in SSR, CSR and SSG architectures, performance optimization, "
        "and accessibility. Familiar with backend technologies including Node.js, Express, MySQL, MongoDB, PostgreSQL, Supabase, Python, and Flask. "
        "Passionate about continuous learning and building seamless user experiences."
    )

    # Skills Bölümü
    cv.add_heading('Skills', level=1)
    skills_list = [
        "HTML5",
        "CSS3",
        "JavaScript (ES6+)",
        "TypeScript",
        "Vue.js 3",
        "Nuxt.js 3",
        "Vue Router",
        "Pinia",
        "SSR, CSR, SSG",
        "Tailwind CSS",
        "SASS/SCSS",
        "RESTful APIs",
        "GraphQL (Apollo Client, graphql)",
        "Vitest, Cypress, Jest",
        "Git, GitHub, VS Code",
        "CI/CD fundamentals",
        "SEO Optimization",
        "Responsive & Mobile-first Design",
    ]

    for skill in skills_list:
        cv.add_paragraph(skill, style='List Bullet')

    # Experience
    cv.add_heading('Experience', level=1)
    exp = cv.add_paragraph()
    exp.add_run("Frontend Developer – Remote | 2022 - Present\n").bold = True
    exp.add_run(
        "- Built and maintained client websites using Vue 3 and Nuxt 3.\n"
        "- Focused on SEO improvements, responsiveness, and accessibility."
    )

    # Projects
    cv.add_heading('Projects', level=1)
    proj = cv.add_paragraph()
    proj.add_run("View all featured projects and case studies at:\n").bold = True
    proj.add_run("https://fuad-foliou.netlify.app\n\n")

    proj.add_run("1. SEO-Optimized Film Platform\n").bold = True
    proj.add_run(
        "A movie discovery platform inspired by Netflix, using TMDB API.\n"
        "Stack: Nuxt.js 3, Tailwind CSS\n"
        "Implemented SSR and GSR to improve SEO.\n"
        "Deployed on Vercel with CI pipelines.\n"
        "https://nuflix-nu.vercel.app\n\n"
    )

    proj.add_run("2. E-Commerce SPA with Vue 3\n").bold = True
    proj.add_run(
        "A modern single-page shop app with cart, product filters and REST API integration.\n"
        "Stack: Vue.js 3, Pinia, Tailwind, REST API\n"
        "Created reusable UI components using Composition API.\n"
        "https://fuad-esite.netlify.app"
    )

    # Education
    cv.add_heading('Education', level=1)
    edu = cv.add_paragraph()
    edu.add_run("Self-Taught Frontend Developer\n").bold = True
    edu.add_run("Completed 3+ years of structured self-learning in modern frontend development (2022–2025).")

    # Languages
    cv.add_heading('Languages', level=1)
    lang = cv.add_paragraph()
    lang.add_run("Azerbaijani: Native\n").bold = True
    lang.add_run("Turkish: Fluent\n").bold = True
    lang.add_run("English: B2 (Intermediate)").bold = True

    cv.save(filename)
    print(f"CV başarıyla oluşturuldu: {filename}")
    return filename

def convert_docx_to_pdf(docx_path, pdf_path):
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(docx_path)
    doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
    doc.Close()
    word.Quit()
    print(f"PDF başarıyla oluşturuldu: {pdf_path}")

if __name__ == "__main__":
    docx_filename = "Fuad_Aghayev_Resume_CV.docx"
    pdf_filename = "Fuad_Aghayev_Resume_CV.pdf"

    docx_path = os.path.abspath(docx_filename)
    pdf_path = os.path.abspath(pdf_filename)

    create_cv_docx(docx_filename)
    convert_docx_to_pdf(docx_path, pdf_path)
